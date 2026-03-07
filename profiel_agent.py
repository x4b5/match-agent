import os
import glob
import json
import re
import requests
from datetime import datetime
from config import (
    OLLAMA_URL,
    PROFIEL_MODEL,
    SYSTEM_PROMPT,
    PROFIEL_KANDIDAAT_PROMPT,
    PROFIEL_WERKGEVERSVRAAG_PROMPT
)


def _vraag_ollama_profiel(prompt: str) -> dict | str:
    """Vraag Ollama om een profiel te genereren. Retourneert dict bij succes, foutmelding (str) bij fout."""
    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": PROFIEL_MODEL,
                "prompt": prompt,
                "system": SYSTEM_PROMPT,
                "stream": False,
                "format": "json",
                "options": {
                    "temperature": 0.1,
                    "num_predict": 1536,
                },
                "think": False,
            },
            timeout=300,
        )
        response.raise_for_status()
    except requests.exceptions.ConnectionError:
        return "Kan geen verbinding maken met Ollama. Draait het? (ollama serve)"
    except requests.exceptions.Timeout:
        return "Timeout bij profiel-extractie (>5 min). Probeer het opnieuw."
    except requests.exceptions.HTTPError as e:
        return f"HTTP fout van Ollama: {e}"
    except Exception as e:
        return f"Onverwachte fout bij profiel-extractie: {e}"

    antwoord = response.json().get("response", "")
    json_match = re.search(r"\{.*\}", antwoord, re.DOTALL)
    if not json_match:
        return "LLM gaf geen geldig JSON-antwoord. Probeer opnieuw."

    try:
        return json.loads(json_match.group())
    except json.JSONDecodeError:
        return "Kon het LLM-antwoord niet als JSON parsen."


def profileer_kandidaat(tekst: str) -> dict | str:
    """Zet ruwe kandidaattekst om naar gestandaardiseerd profiel via LLM."""
    prompt = PROFIEL_KANDIDAAT_PROMPT.format(tekst=tekst)
    return _vraag_ollama_profiel(prompt)


def profileer_werkgeversvraag(tekst: str) -> dict | str:
    """Zet ruwe werkgeversvraag om naar gestandaardiseerd profiel via LLM."""
    prompt = PROFIEL_WERKGEVERSVRAAG_PROMPT.format(tekst=tekst)
    return _vraag_ollama_profiel(prompt)


def vind_meest_recente_profiel(map_pad: str) -> str | None:
    """Vindt het meest recente .json bestand in de map."""
    zoekpad = os.path.join(map_pad, "*.json")
    bestanden = glob.glob(zoekpad)
    if not bestanden:
        return None
    return max(bestanden, key=os.path.getmtime)


def laad_profiel_uit_map(map_pad: str) -> dict | None:
    """Laad het meest recente profiel (.json) uit een specifieke map, indien aanwezig."""
    profiel_pad = vind_meest_recente_profiel(map_pad)
    if profiel_pad and os.path.exists(profiel_pad):
        try:
            with open(profiel_pad, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"Fout bij lezen {profiel_pad}: {e}")
    return None


def genereer_profiel_voor_map(map_pad: str, profileer_fn) -> dict | str:
    """
    Lees alle .txt en .docx bestanden in de map, plak ze aan elkaar,
    stuur naar Ollama, en sla het resultaat op als <mapnaam>_<tijdstempel>.json in de map.

    Retourneert het profiel (dict) bij succes, of een foutmelding (str) bij fout.
    """
    if not os.path.isdir(map_pad):
        return f"{map_pad} is geen geldige map."

    gecombineerde_tekst = ""
    baseless_bestanden = os.listdir(map_pad)
    tekst_extensies = (".txt", ".md", ".csv", ".eml", ".json", ".log", ".rtf")
    verwerkbare_bestanden = [f for f in baseless_bestanden if f.lower().endswith(tekst_extensies) or f.lower().endswith(".docx") or f.lower().endswith(".pdf")]

    if not verwerkbare_bestanden:
        return "Geen leesbare bestanden gevonden in de map."

    waarschuwingen = []

    for doc in sorted(verwerkbare_bestanden):
        pad = os.path.join(map_pad, doc)
        gecombineerde_tekst += f"\n\n--- Inhoud uit {doc} ---\n\n"
        try:
            if doc.lower().endswith(tekst_extensies):
                with open(pad, "r", encoding="utf-8") as f:
                    gecombineerde_tekst += f.read()
            elif doc.endswith(".docx"):
                try:
                    from docx import Document
                    document = Document(pad)
                    voor_docx = "\n".join([p.text for p in document.paragraphs])
                    gecombineerde_tekst += voor_docx
                except ImportError:
                    waarschuwingen.append(f"python-docx module ontbreekt — {doc} overgeslagen.")
                except Exception as e:
                    waarschuwingen.append(f"Kon {doc} niet lezen: {e}")
            elif doc.endswith(".pdf"):
                try:
                    import pypdf
                    reader = pypdf.PdfReader(pad)
                    voor_pdf = ""
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            voor_pdf += text + "\n"
                    gecombineerde_tekst += voor_pdf
                except ImportError:
                    waarschuwingen.append(f"pypdf module ontbreekt — {doc} overgeslagen.")
                except Exception as e:
                    waarschuwingen.append(f"Kon {doc} niet lezen: {e}")
        except Exception as e:
            waarschuwingen.append(f"Kon {doc} niet lezen: {e}")

    if not gecombineerde_tekst.strip():
        return "Geen leesbare tekst gevonden in de bestanden."

    resultaat = profileer_fn(gecombineerde_tekst)

    # Als profileer_fn een foutmelding retourneert (str), geef die door
    if isinstance(resultaat, str):
        return resultaat

    if resultaat and isinstance(resultaat, dict):
        # Sla waarschuwingen op in het profiel als metadata
        if waarschuwingen:
            resultaat["_waarschuwingen"] = waarschuwingen

        map_naam = os.path.basename(os.path.normpath(map_pad))
        tijdstempel = datetime.now().strftime("%Y%m%d_%H%M%S")
        profiel_naam = f"{map_naam}_{tijdstempel}.json"
        profiel_pad = os.path.join(map_pad, profiel_naam)
        try:
            with open(profiel_pad, "w", encoding="utf-8") as f:
                json.dump(resultaat, f, ensure_ascii=False, indent=2)
            return resultaat
        except Exception as e:
            return f"Kon profiel niet opslaan: {e}"

    return "LLM gaf een leeg antwoord."

def laad_ruwe_tekst_uit_map(map_pad: str) -> str:
    """Verzamelt alle leesbare tekst uit de map en combineert dit."""
    if not os.path.isdir(map_pad):
        return ""

    gecombineerde_tekst = ""
    baseless_bestanden = os.listdir(map_pad)
    tekst_extensies = (".txt", ".md", ".csv", ".eml", ".json", ".log", ".rtf")
    verwerkbare_bestanden = [f for f in baseless_bestanden if f.lower().endswith(tekst_extensies) or f.lower().endswith(".docx") or f.lower().endswith(".pdf")]

    for doc in sorted(verwerkbare_bestanden):
        pad = os.path.join(map_pad, doc)
        gecombineerde_tekst += f"\n\n--- Inhoud uit {doc} ---\n\n"
        try:
            if doc.lower().endswith(tekst_extensies):
                with open(pad, "r", encoding="utf-8") as f:
                    gecombineerde_tekst += f.read()
            elif doc.endswith(".docx"):
                try:
                    from docx import Document
                    document = Document(pad)
                    voor_docx = "\n".join([p.text for p in document.paragraphs])
                    gecombineerde_tekst += voor_docx
                except Exception:
                    pass
            elif doc.endswith(".pdf"):
                try:
                    import pypdf
                    reader = pypdf.PdfReader(pad)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            gecombineerde_tekst += text + "\n"
                except Exception:
                    pass
        except Exception:
            pass
    return gecombineerde_tekst.strip()
